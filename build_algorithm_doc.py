from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"F:\计图")
OUT_DIR = ROOT / "交付"
OUT_PATH = OUT_DIR / "计图挑战赛_代码复现与核心算法说明.docx"

# compact_reference_guide preset, with one named override:
# Latin font Calibri; East Asian font Microsoft YaHei for reliable Chinese rendering.
INK = "172033"
NAVY = "1F4D78"
BLUE = "2E74B5"
MUTED = "5E6B7A"
LIGHT_BLUE = "E8EEF5"
LIGHTER_BLUE = "F4F7FA"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "D9E0E7"
WHITE = "FFFFFF"
CAUTION = "7A5A00"
CAUTION_FILL = "FFF8E8"
RISK = "9B1C1C"
RISK_FILL = "FDF1F1"
POSITIVE = "1F3A5F"
POSITIVE_FILL = "EEF5FB"

LATIN = "Calibri"
CJK = "Microsoft YaHei"
MONO = "Consolas"


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_run_font(run, *, latin=LATIN, east_asia=CJK, size=None, color=None,
                 bold=None, italic=None):
    run.font.name = latin
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:eastAsia"), east_asia)
    rfonts.set(qn("w:cs"), latin)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_repeat_table_header(row):
    trpr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    trpr.append(tbl_header)


def set_cell_shading(cell, fill: str):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcpr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tcpr = cell._tc.get_or_add_tcPr()
    tc_mar = tcpr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tcpr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int):
    tcpr = cell._tc.get_or_add_tcPr()
    tcw = tcpr.find(qn("w:tcW"))
    if tcw is None:
        tcw = OxmlElement("w:tcW")
        tcpr.append(tcw)
    tcw.set(qn("w:w"), str(width_dxa))
    tcw.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tblpr = table._tbl.tblPr
    tblw = tblpr.find(qn("w:tblW"))
    if tblw is None:
        tblw = OxmlElement("w:tblW")
        tblpr.append(tblw)
    tblw.set(qn("w:w"), str(total))
    tblw.set(qn("w:type"), "dxa")
    tblind = tblpr.find(qn("w:tblInd"))
    if tblind is None:
        tblind = OxmlElement("w:tblInd")
        tblpr.append(tblind)
    tblind.set(qn("w:w"), str(indent_dxa))
    tblind.set(qn("w:type"), "dxa")
    layout = tblpr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tblpr.append(layout)
    layout.set(qn("w:type"), "fixed")

    old_grid = table._tbl.tblGrid
    for child in list(old_grid):
        old_grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        old_grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            set_cell_width(cell, width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table, color=MID_GRAY, size=5):
    tblpr = table._tbl.tblPr
    borders = tblpr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblpr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def add_page_field(paragraph):
    run = paragraph.add_run("第 ")
    set_run_font(run, size=9, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    r = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), LATIN)
    rfonts.set(qn("w:hAnsi"), LATIN)
    rfonts.set(qn("w:eastAsia"), CJK)
    rpr.append(rfonts)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), MUTED)
    rpr.append(color)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "18")
    rpr.append(sz)
    r.append(rpr)
    t = OxmlElement("w:t")
    t.text = "1"
    r.append(t)
    fld.append(r)
    paragraph._p.append(fld)
    run = paragraph.add_run(" 页")
    set_run_font(run, size=9, color=MUTED)


def paragraph_shading(paragraph, fill: str):
    ppr = paragraph._p.get_or_add_pPr()
    shd = ppr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        ppr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def paragraph_left_border(paragraph, color: str, size=16, space=8):
    ppr = paragraph._p.get_or_add_pPr()
    pbdr = ppr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        ppr.append(pbdr)
    left = pbdr.find(qn("w:left"))
    if left is None:
        left = OxmlElement("w:left")
        pbdr.append(left)
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(size))
    left.set(qn("w:space"), str(space))
    left.set(qn("w:color"), color)


def configure_numbering(doc: Document):
    # Use Word's built-in list styles without altering numbering.xml. The two
    # numbered sections use different built-in styles, so each remains a proper
    # semantic list while restarting independently.
    return "List Bullet", "List Number", "List Number 2"


def apply_num(paragraph, num_id: int):
    ppr = paragraph._p.get_or_add_pPr()
    numpr = ppr.find(qn("w:numPr"))
    if numpr is None:
        numpr = OxmlElement("w:numPr")
        ppr.append(numpr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    nid = OxmlElement("w:numId")
    nid.set(qn("w:val"), str(num_id))
    numpr.extend([ilvl, nid])


def setup_styles(doc: Document):
    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = LATIN
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal._element.rPr.rFonts.set(qn("w:ascii"), LATIN)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), LATIN)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), CJK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, NAVY, 10, 5),
    ):
        style = styles[name]
        style.font.name = LATIN
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style._element.rPr.rFonts.set(qn("w:ascii"), LATIN)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), LATIN)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), CJK)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.widow_control = True

    if "Doc Caption" not in styles:
        cap = styles.add_style("Doc Caption", WD_STYLE_TYPE.PARAGRAPH)
    else:
        cap = styles["Doc Caption"]
    cap.font.name = LATIN
    cap.font.size = Pt(9)
    cap.font.italic = True
    cap.font.color.rgb = rgb(MUTED)
    cap._element.rPr.rFonts.set(qn("w:ascii"), LATIN)
    cap._element.rPr.rFonts.set(qn("w:hAnsi"), LATIN)
    cap._element.rPr.rFonts.set(qn("w:eastAsia"), CJK)
    cap.paragraph_format.space_before = Pt(4)
    cap.paragraph_format.space_after = Pt(4)
    cap.paragraph_format.keep_with_next = True

    if "Code Block" not in styles:
        code = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code = styles["Code Block"]
    code.font.name = MONO
    code.font.size = Pt(8.5)
    code.font.color.rgb = rgb(INK)
    code._element.rPr.rFonts.set(qn("w:ascii"), MONO)
    code._element.rPr.rFonts.set(qn("w:hAnsi"), MONO)
    code._element.rPr.rFonts.set(qn("w:eastAsia"), CJK)
    code.paragraph_format.left_indent = Inches(0.18)
    code.paragraph_format.right_indent = Inches(0.12)
    code.paragraph_format.space_before = Pt(4)
    code.paragraph_format.space_after = Pt(6)
    code.paragraph_format.line_spacing = 1.05
    code.paragraph_format.widow_control = True


def configure_section(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def set_running_furniture(doc: Document):
    for section in doc.sections:
        configure_section(section)
        header = section.header
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run("计图挑战赛 · 代码复现与核心算法")
        set_run_font(run, size=8.5, color=MUTED, bold=False)
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.paragraph_format.space_before = Pt(0)
        fp.paragraph_format.space_after = Pt(0)
        add_page_field(fp)


def add_body(doc, text, *, bold_prefix=None, after=None, keep=False):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_run_font(r, bold=True, color=INK)
        r = p.add_run(text[len(bold_prefix):])
        set_run_font(r, color=INK)
    else:
        r = p.add_run(text)
        set_run_font(r, color=INK)
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.keep_with_next = keep
    return p


def add_bullet(doc, text, bullet_id, *, bold_prefix=None):
    p = doc.add_paragraph(style=bullet_id)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_run_font(r, bold=True)
        r = p.add_run(text[len(bold_prefix):])
        set_run_font(r)
    else:
        r = p.add_run(text)
        set_run_font(r)
    return p


def add_numbered(doc, text, number_id, *, bold_prefix=None):
    p = doc.add_paragraph(style=number_id)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_run_font(r, bold=True)
        r = p.add_run(text[len(bold_prefix):])
        set_run_font(r)
    else:
        r = p.add_run(text)
        set_run_font(r)
    return p


def add_callout(doc, label, text, *, kind="info"):
    if kind == "risk":
        fill, accent = RISK_FILL, RISK
    elif kind == "caution":
        fill, accent = CAUTION_FILL, CAUTION
    elif kind == "positive":
        fill, accent = POSITIVE_FILL, POSITIVE
    else:
        fill, accent = LIGHTER_BLUE, BLUE
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.10)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(8)
    paragraph_shading(p, fill)
    paragraph_left_border(p, accent)
    r = p.add_run(label + "  ")
    set_run_font(r, bold=True, color=accent)
    r = p.add_run(text)
    set_run_font(r, color=INK)
    return p


def add_code(doc, text):
    p = doc.add_paragraph(style="Code Block")
    paragraph_shading(p, LIGHT_GRAY)
    for idx, line in enumerate(text.splitlines()):
        if idx:
            p.add_run().add_break()
        r = p.add_run(line)
        set_run_font(r, latin=MONO, east_asia=CJK, size=8.5, color=INK)
    return p


def add_equation(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(7)
    r = p.add_run(text)
    set_run_font(r, latin="Cambria Math", east_asia=CJK, size=10.5, color=NAVY, italic=True)
    return p


def add_table(doc, headers, rows, widths_dxa, *, font_size=9.2, header_fill=LIGHT_BLUE,
              first_col_bold=False, alignments=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths_dxa, indent_dxa=120)
    set_table_borders(table)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for idx, label in enumerate(headers):
        cell = hdr.cells[idx]
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.05
        r = p.add_run(str(label))
        set_run_font(r, size=font_size, bold=True, color=NAVY)
    for ridx, row in enumerate(rows):
        cells = table.add_row().cells
        for cidx, value in enumerate(row):
            if ridx % 2 == 1:
                set_cell_shading(cells[cidx], "FAFBFC")
            p = cells[cidx].paragraphs[0]
            if alignments and cidx < len(alignments):
                p.alignment = alignments[cidx]
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.10
            r = p.add_run(str(value))
            set_run_font(r, size=font_size, bold=(first_col_bold and cidx == 0), color=INK)
    set_table_geometry(table, widths_dxa, indent_dxa=120)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_cover(doc: Document):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(70)
    p.paragraph_format.space_after = Pt(18)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("算法复现报告 · 正式赛道 Track 1")
    set_run_font(r, size=11, color=BLUE, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("计图挑战赛代码复现\n与核心算法说明")
    set_run_font(r, size=28, color=NAVY, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(30)
    r = p.add_run("时间动态图候选边排序、Learning-to-Rank、图 SVD 与融合优化")
    set_run_font(r, size=14, color=MUTED)

    add_callout(
        doc,
        "复现结论",
        "正式赛最终主线不是单一 GNN，而是以时间有向图统计为底座，把每个测试样本的 100 个候选节点组成一个排序组，使用 XGBoost Ranker 学习候选相对次序，再按数据集采用不同特征配置和后处理融合。",
        kind="positive",
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(44)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("依据：工作区现有源码、版本日志与 metadata.json")
    set_run_font(r, size=10, color=MUTED, italic=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("整理日期：2026-07-30")
    set_run_font(r, size=10, color=MUTED)


def build_document():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    setup_styles(doc)
    for section in doc.sections:
        configure_section(section)
    bullet_id, number_id, decision_number_id = configure_numbering(doc)

    props = doc.core_properties
    props.title = "计图挑战赛代码复现与核心算法说明"
    props.subject = "时间动态图候选边排序、XGBoost Learning-to-Rank、SVD 与融合"
    props.author = "Codex（依据工作区代码整理）"
    props.keywords = "计图挑战赛, 图学习, 时间动态图, Learning-to-Rank, XGBoost, SVD, GCN"
    props.comments = "基于 F:\\计图 工作区源码和训练日志生成。"

    add_cover(doc)
    doc.add_page_break()

    doc.add_heading("1. 结论摘要", level=1)
    add_body(
        doc,
        "从代码接口可以还原出正式赛任务：给定源节点 src、时间 time，以及 c1～c100 共 100 个候选目标节点，需要为每个候选输出一个分数，使真实目标尽可能排在前面。训练数据是带时间戳的有向边 (src, dst, time)。因此本质是“时间动态图上的候选边排序”，而不是普通节点分类。",
    )
    add_callout(
        doc,
        "一句话算法",
        "先把历史边压缩成可向量化查询的计数、最近时间、转移、跳跃转移与图嵌入索引；再对每行 100 个候选构造绝对特征、行内相对排名特征和测试候选分布特征；最后用 XGBRanker 以 NDCG 目标训练，并用行内归一化及可选的秩融合生成提交。",
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("原始 CSV  →  历史图索引  →  180 维候选特征  →  XGBRanker  →  行内归一化  →  分数据集组装/融合")
    set_run_font(r, size=11, bold=True, color=NAVY)

    add_table(
        doc,
        ["阶段", "dataset1 本地 MRR", "dataset2 本地 MRR", "关键变化"],
        [
            ("V18 启发式", "0.672611", "0.419911", "历史对、时序转移、跳跃转移、热度的手工加权"),
            ("V58 排序器", "—", "0.484296", "130k 排序组、850 树、NDCG、测试分布特征"),
            ("V59 图 SVD", "—", "0.549527", "加入 32 维有向图低秩结构，形成最大单步增益"),
            ("V62 双数据集 Ranker", "0.840013", "0.549600", "dataset1 也切换至 v61 排序器"),
            ("V69 最终独立方案", "0.845617", "0.549600（复用 V62）", "dataset1 用 v65 特征、1200 树；总和 1.395217"),
        ],
        [1900, 1500, 1650, 4310],
        font_size=8.8,
        first_col_bold=True,
        alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT],
    )
    add_body(
        doc,
        "表 1 只记录代码中的本地验证结果。不同版本的采样行数、树数或候选构造方式并不完全相同，不能把所有差值都当作严格消融；其中 V58→V59、V59→V61 的同协议对比最有解释力。",
        after=8,
    )

    doc.add_heading("最终建议复现配置", level=2)
    add_bullet(doc, "dataset1：v65 特征子集，去掉候选位置特征，32 维正向/反向 SVD，row 候选采样，XGBRanker 1200 棵树，最终 ranker 权重 α=0.95。", bullet_id, bold_prefix="dataset1：")
    add_bullet(doc, "dataset2：v61 精简特征子集，去掉候选位置特征，32 维正向/反向 SVD，row 候选采样，XGBRanker 850 棵树，ranker 权重 α=1.00。", bullet_id, bold_prefix="dataset2：")
    add_bullet(doc, "可选后处理：与外部分数做 reciprocal-rank、RRF、Borda、归一化分数或置信度门控融合；若只要求完全自包含复现，应以 V69 独立方案为主。", bullet_id, bold_prefix="可选后处理：")

    doc.add_heading("2. 数据、接口与评价目标", level=1)
    doc.add_heading("2.1 正式赛数据结构", level=2)
    add_table(
        doc,
        ["数据集", "train.csv", "test.csv", "训练字段", "测试字段"],
        [
            ("dataset1", "690,848 行", "61,051 行", "src, dst, time", "src, time, c1…c100"),
            ("dataset2", "2,261,283 行", "153,420 行", "src, dst, time, split", "src, time, c1…c100"),
        ],
        [1250, 1400, 1400, 2350, 2960],
        font_size=9.2,
        first_col_bold=True,
        alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
    )
    add_body(doc, "dataset2 的 split=0 有 1,995,088 行，split=1 有 266,195 行；代码把 split=0 当作历史，把 split=1 当作时间后移验证集。dataset1 没有 split 字段，因此按时间排序后使用前 85% 作为历史、后 15% 作为验证。")
    add_body(doc, "提交文件 dataset1.csv 与 dataset2.csv 每行必须有 100 个浮点数，列顺序对应 c1～c100。代码输出的是行内 min-max 归一化分数，不是经过校准的概率。")

    doc.add_heading("2.2 本地排序指标", level=2)
    add_body(doc, "验证集先把真实 dst 随机插入某个候选位置，再计算真实候选的名次。MRR（Mean Reciprocal Rank）强调把真实目标放在最前面：")
    add_equation(doc, "MRR = (1 / N) · Σᵢ 1 / rankᵢ")
    add_body(doc, "XGBRanker 训练使用 rank:ndcg，监控 ndcg@10；最终选择仍按 MRR 比较。每个测试行是一个 query group，组大小固定为 100，标签只有一个正例。")

    doc.add_heading("2.3 候选构造为什么是关键", level=2)
    add_body(doc, "早期验证直接从所有候选平铺池随机负采样（flat），与正式测试“整行 100 候选共同出现”的分布不一致。后续版本改为 row：随机抽取一整行官方候选，再将真实目标插入随机位置，并移除重复真值。这一改变让排序器学习到真实候选集中的相对难度，而不是过于容易的随机负样本。")
    add_callout(doc, "复现要点", "最终配置应同时设置训练候选模式和评估候选模式为 row；否则即使特征和树参数相同，本地 MRR 也不可直接对照。", kind="caution")

    doc.add_heading("3. 代码结构与执行链路", level=1)
    add_table(
        doc,
        ["文件", "角色", "核心入口/函数"],
        [
            ("gcn.py", "热身赛 Cora 节点分类", "build_norm_adj, GCN, select_epoch, quota_assign"),
            ("track1_solution_v62_dataset1_ranker.py", "正式赛主实现；当前文件已包含 v65/v67 后续特征配置", "build_history, feature_batch, train_ranker, solve_dataset"),
            ("fuse_submissions.py", "两份提交的 reciprocal-rank 融合", "reciprocal_rank_scores, normalize_rows"),
            ("fusion_experiments_v71.py", "RRF/Borda/rank-product/自适应融合实验", "method_scores, confidence"),
            ("gated_fusions_v88.py", "按行置信度选择性切换模型", "gated_dataset1"),
            ("hybrid_copy_fusions_v83.py", "按数据集直接选择来源并重新打包", "HYBRIDS, read_bytes"),
        ],
        [2900, 2850, 3610],
        font_size=8.7,
        first_col_bold=True,
    )

    doc.add_heading("主程序执行顺序", level=2)
    for text in [
        "读取并按 time 稳定排序训练边；必要时下载并解压官方数据。",
        "构建历史索引：节点统计、(src,dst) 对统计、顺序转移、1～5 跳跳跃转移、最近 10 个目标、正向/反向 SVD，以及测试候选分布统计。",
        "按验证规则生成伪测试行；对每行 100 个候选批量构造特征。",
        "训练 XGBRanker；GPU 失败时自动回退到 CPU hist。",
        "在 α∈[0,1] 的网格上比较 ranker 分数与 V18 启发式基线的混合 MRR。",
        "在完整历史和正式 test 上重建索引，分块写出两个 CSV，再压缩为 result.zip。",
    ]:
        add_numbered(doc, text, decision_number_id)

    doc.add_heading("4. 热身赛：两层 GCN 的复现逻辑", level=1)
    add_body(doc, "热身赛代码处理 Cora 节点分类。虽然它不是正式赛最终主线，但体现了图卷积、重复训练与结构化后处理的基本思路。")

    doc.add_heading("4.1 图归一化与模型", level=2)
    add_body(doc, "先对节点词袋特征做行归一化，再给邻接矩阵加自环，并做对称归一化：")
    add_equation(doc, "X̄ᵢ = Xᵢ / (Σⱼ Xᵢⱼ + ε),     Â = D⁻¹ᐟ² (A + I) D⁻¹ᐟ²")
    add_body(doc, "代码提前计算 AX=ÂX，随后执行两层传播：")
    add_equation(doc, "H = Dropout(ReLU((ÂX)W₀)),     Z = (ÂH)W₁")
    add_body(doc, "这里第一层的邻接传播通过缓存 ÂX 完成，第二层用 torch.sparse.mm 再传播一次；隐藏维 256，dropout 0.75，无偏置线性层，Adam 学习率 0.01、weight decay 5×10⁻⁴。")

    doc.add_heading("4.2 训练稳定性优化", level=2)
    add_bullet(doc, "早停选轮数：每个随机种子只在 train_mask 上训练，用 val_mask 选择最佳 epoch，并以 patience=80 提前终止。", bullet_id, bold_prefix="早停选轮数：")
    add_bullet(doc, "train+val 重训：若 final_train_mask=train_val，则用选出的最佳轮数从头重训，把验证标签也用于最终拟合。", bullet_id, bold_prefix="train+val 重训：")
    add_bullet(doc, "多随机种子集成：默认 8 次运行，种子按 seed+97×run_id 变化，最后平均 softmax 概率，降低单次训练方差。", bullet_id, bold_prefix="多随机种子集成：")
    add_bullet(doc, "设备降级：CUDA 不可用或计算能力过低时回退 CPU，保证脚本可运行。", bullet_id, bold_prefix="设备降级：")

    doc.add_heading("4.3 类别配额后处理", level=2)
    add_body(doc, "代码利用公开聚合反馈给测试集设定 7 类目标数量 [105,151,129,148,310,92,65]。它把每个类别展开成对应数量的“槽位”，然后求解：")
    add_equation(doc, "min Σᵢ −log P(i, yᵢ),   s.t. 每类预测数量严格等于 quota")
    add_body(doc, "若 SciPy 可用，使用 Hungarian/线性分配得到全局最优标签；否则从超额类别向不足类别做贪心迁移。这个优化不是模型学习，而是把已知的类别先验作为全局约束。")
    add_callout(doc, "风险提示", "类别配额强依赖外部反馈。若测试分布或评分轮次改变，固定 quota 可能造成反效果；复现时应把它作为可开关的后处理，而不是 GCN 的必需组成。", kind="risk")

    doc.add_heading("5. 正式赛核心算法：时间动态图候选排序", level=1)
    doc.add_heading("5.1 历史索引：把图查询变成数组查询", level=2)
    add_body(doc, "核心工程思想是：昂贵的 pandas groupby 只做一次，预测阶段不逐行查字典。所有二元关系都编码为 key = u × base + v，排序后用 np.searchsorted 批量查找；节点级统计放入稠密数组，候选节点 ID 可直接索引。")
    add_table(
        doc,
        ["索引/统计", "定义", "表达的信息"],
        [
            ("pair(src,dst)", "出现次数、最后时间", "源节点是否曾直接连接候选；强记忆特征"),
            ("transition(prev,dst)", "同一 src 的相邻目标转移", "最近一次访问目标之后，下一目标通常是谁"),
            ("skip transition", "同一 src 内间隔 1～5 的目标转移", "比一阶转移更鲁棒的短序列模式"),
            ("dst/src node stats", "全局次数、最近时间、多个时间窗口次数", "候选热度、近期趋势、节点活跃度"),
            ("recent_by_src", "每个 src 最近 10 个 dst", "短期兴趣与局部上下文"),
            ("test candidate stats", "候选在 test 中的频率、首末出现时间、src-candidate 共现", "无标签的测试分布先验"),
            ("directed SVD", "正向与反向图的低秩嵌入", "高阶结构、角色相似度和方向性"),
        ],
        [1900, 3050, 4410],
        font_size=8.8,
        first_col_bold=True,
    )

    doc.add_heading("5.2 V18 启发式基线", level=2)
    add_body(doc, "对第 i 行的候选 d，启发式分数是多组证据的线性和。计数特征用 log1p 压缩长尾，最近时间用指数衰减：")
    add_equation(doc, "recency(t,T) = exp(−max(t−T,0)/τ),     τ = (tₘₐₓ−tₘᵢₙ)/10")
    add_equation(doc, "S(s,d,t)=wₚ log(1+Nₛd)+wᵣ Rₛd + w_d·Pop(d)+w_t·TestFreq(d) + Σⱼ[wⱼᴹ Matchⱼ + wⱼᵀ Transⱼ + wⱼˢ Skipⱼ] + ε")
    add_body(doc, "dataset1 更依赖 pair 与时序转移，且显著下调全局热度；dataset2 把直接 pair 权重设为 0，显著提高近期目的节点热度、测试候选频率与 skip transition。这说明两个数据集的生成机制不同，统一一套权重会损失效果。")

    doc.add_heading("5.3 180 维原始特征体系", level=2)
    add_table(
        doc,
        ["特征组", "代表特征", "为什么有效"],
        [
            ("直接/反向边", "log_pair_cnt, pair_rec, log_rev_pair_cnt", "记忆已有关系，并区分有向图两个方向"),
            ("节点热度与窗口", "log_dst_window_50…99, dst_rec", "把长期热度与最近爆发分开"),
            ("测试分布", "log_test_freq, test_src_cand_freq, first/last/span", "匹配正式候选生成分布"),
            ("行内相对特征", "row_norm, inverse_rank, zscore, margin_best", "排序只关心同一行候选的相对强弱"),
            ("最近历史", "recent_match_0…9, recent_common_*", "捕捉短期重复、共同近期邻居与回访"),
            ("转移/跳转", "trans/skip 及其 reverse、recency", "将时间序列模式转成可学习特征"),
            ("图低秩结构", "svd_dot, svd_cos, svd_inv_rank", "弥补局部计数对未见边和高阶关系的不足"),
            ("V18 教师特征", "v18_score/norm/inv_rank/zscore", "让树模型从强手工基线起步，再学习非线性修正"),
            ("候选位置", "pos_norm, inv_pos", "容易利用列顺序伪相关；最终默认移除"),
        ],
        [1850, 3310, 4200],
        font_size=8.7,
        first_col_bold=True,
    )

    doc.add_heading("5.4 Learning-to-Rank", level=2)
    add_body(doc, "对每个验证样本构造 100×F 特征矩阵，真实候选标签为 1，其余为 0，group=[100,100,…]。XGBRanker 直接优化组内排序，能够学习“热度在冷门 src 上是否该降权”“SVD 相似度与 pair 计数如何交互”等非线性规则。")
    add_body(doc, "训练后分别把 ranker 原始分数和 V18 分数做行内 min-max 归一化，再搜索混合系数 α：")
    add_equation(doc, "S_final = α · Norm(S_ranker) + (1−α) · Norm(S_V18)")
    add_body(doc, "最终 dataset2 选择 α=1.00，说明排序器完全替代启发式；dataset1 的 V69 选择 α=0.95，保留 5% 基线作为轻微稳定器。")

    doc.add_heading("5.5 有向图 SVD：最大单步增益", level=2)
    add_body(doc, "将训练边构造成稀疏有向矩阵 M。每条边的权重随时间线性增加，最近边最高为 1.35 倍：")
    add_equation(doc, "M_uv = Σ_(u→v,t) [1 + 0.35·(t−tₘᵢₙ)/(tₘₐₓ−tₘᵢₙ)]")
    add_body(doc, "TruncatedSVD 得到 32 维源嵌入和目标嵌入，候选特征包括点积、余弦、行内归一化、倒数名次和向量范数；另对反向图再做一次 SVD，保留方向角色。V58→V59 在同为 130k 组、850 树、NDCG 的记录中，dataset2 MRR 从 0.484296 提升到 0.549527，是整条优化链最明显的提升。")
    add_callout(doc, "解释", "局部计数擅长“见过的边”和短序列模式；SVD 能把共同出边/入边结构压缩成连续空间，对未直接出现的候选关系提供高阶泛化。", kind="positive")

    doc.add_heading("5.6 测试分布特征的性质", level=2)
    add_body(doc, "代码会使用整个无标签 test 的候选频率、候选首次/最后出现时间以及 src-candidate 在 test 中的共现频率。这属于 transductive 特征：没有读取真实标签，但利用了测试输入的整体分布。伪验证也用相同方式在 fake_test 上构建这些统计，从而尽量保持训练与推理一致。")
    add_callout(doc, "规则检查", "若比赛规则禁止跨测试样本聚合，必须关闭 test_cand_* 和 test_src_cand_* 特征；如果允许使用完整测试输入，则这是方案的重要组成。", kind="caution")

    doc.add_heading("6. 优化路线复盘", level=1)
    add_table(
        doc,
        ["版本阶段", "主要变化", "验证结论/作用"],
        [
            ("V3–V18", "手工加权：pair、recency、转移、skip、热度；按数据集调权", "得到强基线：d1=0.672611，d2=0.419911"),
            ("V25–V41", "引入 XGBRanker；增加反向边、窗口、两跳/图 motif 等特征", "学习非线性交互，d2 进入 0.46 左右"),
            ("V42", "特征筛选；比较含位置、去位置、全特征", "去位置 0.462936，高于含位置 0.462497 和全特征 0.462530"),
            ("V47–V50", "排序组扩到 130k，树数 700→850", "d2 从 0.476948 到 0.479150"),
            ("V51–V58", "候选按整行采样；加入 test 时间窗口与首末/跨度特征；NDCG", "d2 达到 0.484296"),
            ("V59", "32 维有向 SVD", "d2=0.549527，最显著的结构特征增益"),
            ("V60–V61", "反向 SVD，再裁剪低收益反向特征", "80k/500 同协议下 v61=0.540142，略高于 v60=0.539775"),
            ("V62", "把 v61 ranker 扩展到 dataset1", "d1=0.840013；d2=0.549600"),
            ("V65–V69", "d1 加最近共同邻居、完整正反转移；v65 167 维；1200 树", "d1=0.845617，最终独立总和=1.395217"),
            ("V71–V94", "RR/RRF/Borda/分数融合、按数据集复制、置信度门控", "利用外部模型互补；不再是纯自包含方案"),
        ],
        [1500, 4450, 3410],
        font_size=8.4,
        first_col_bold=True,
    )

    doc.add_heading("最重要的五个优化判断", level=2)
    for text in [
        "先把验证候选做得像测试候选。row 采样减少负采样分布错位，比盲目堆特征更基础。",
        "排序问题优先构造行内相对特征。相同绝对热度在不同候选集合中意义不同，inverse rank、row norm、z-score 比原值更稳定。",
        "删除候选位置特征。真实答案被随机插入验证行，位置本应无信息；去掉它降低了对列顺序偶然性的依赖。",
        "用低秩图结构补充局部统计。SVD 是 dataset2 最大的单步提升来源，且正向/反向分解符合有向图角色差异。",
        "分数据集建模。dataset1 偏记忆与序列，dataset2 偏全局热度、图结构和测试分布；最终分别选择 v65 与 v61。",
    ]:
        add_numbered(doc, text, number_id)

    doc.add_heading("7. 最终独立方案的精确配置", level=1)
    add_table(
        doc,
        ["参数", "dataset1（V69）", "dataset2（V62）"],
        [
            ("feature_profile", "v65：167/180，无位置", "v61：105/180，无位置"),
            ("candidate_mode / eval", "row / row", "row / row"),
            ("ranker_rows", "130000 上限（实际验证可更少）", "130000"),
            ("objective / metric", "rank:ndcg / ndcg@10", "rank:ndcg / ndcg@10"),
            ("trees / depth", "1200 / 5", "850 / 5"),
            ("learning_rate", "0.032", "0.032"),
            ("subsample / colsample", "0.88 / 0.90", "0.88 / 0.90"),
            ("min_child_weight", "20", "20"),
            ("reg_lambda / reg_alpha", "4.0 / 0.0", "4.0 / 0.0"),
            ("SVD", "dim=32, n_iter=5，正向+反向", "dim=32, n_iter=5，正向+反向"),
            ("最终 α", "0.95", "1.00"),
            ("随机种子", "20260525", "20260525"),
        ],
        [2650, 3355, 3355],
        font_size=8.8,
        first_col_bold=True,
    )

    doc.add_heading("7.1 环境准备", level=2)
    add_body(doc, "建议 Python 3.10+，依赖 numpy、pandas、scipy、scikit-learn、xgboost。XGBoost 2.x 使用 tree_method=hist, device=cuda；旧版使用 gpu_hist。没有 GPU 时脚本自动回退 CPU hist，但 130k 排序组会明显更慢。")
    add_code(doc, "python -m pip install numpy pandas scipy scikit-learn xgboost")

    doc.add_heading("7.2 两阶段复现 V69", level=2)
    add_body(doc, "主脚本在最后总是尝试同时打包 dataset1.csv 与 dataset2.csv。最稳妥的方式是先在同一个输出目录生成 V62 两个数据集，再只重训 dataset1 覆盖旧文件；这样第二次运行打包时 dataset2.csv 已存在。")
    add_code(
        doc,
        "$SOL = '.\\formal_track1\\track1_solution_v62_dataset1_ranker.py'\n"
        "$DATA = '.\\formal_track1\\repro_data'\n"
        "$ARCHIVE = '.\\formal_track1\\kaggle_output\\track1_data.zip'\n"
        "$OUT = '.\\formal_track1\\repro_v69'\n"
        "$ZIP = '.\\formal_track1\\result_reproduced_v69.zip'",
    )
    add_body(doc, "阶段 A：按 v61、850 树生成 V62 两个数据集，保留其中 dataset2.csv。")
    add_code(
        doc,
        "python $SOL --data_dir $DATA --zip_path $ARCHIVE --output_dir $OUT `\n"
        "  --result_zip '.\\formal_track1\\result_stage_a_v62.zip' `\n"
        "  --ranker_dataset both --only_dataset both `\n"
        "  --ranker_feature_profile v61 --ranker_candidate_mode row `\n"
        "  --ranker_eval_candidate_mode row --ranker_rows 130000 `\n"
        "  --ranker_trees 850 --ranker_objective ndcg --ranker_eval_metric ndcg@10 `\n"
        "  --ranker_svd_dim 32 --ranker_no_pos",
    )
    add_body(doc, "阶段 B：在同一输出目录只重训 dataset1，使用 v65、1200 树与 α=0.95；主程序会把新 dataset1.csv 和阶段 A 的 dataset2.csv 一起打包。")
    add_code(
        doc,
        "python $SOL --data_dir $DATA --zip_path $ARCHIVE --output_dir $OUT `\n"
        "  --result_zip $ZIP --ranker_dataset dataset1 --only_dataset dataset1 `\n"
        "  --ranker_feature_profile v65 --ranker_candidate_mode row `\n"
        "  --ranker_eval_candidate_mode row --ranker_rows 130000 `\n"
        "  --ranker_trees 1200 --ranker_objective ndcg --ranker_eval_metric ndcg@10 `\n"
        "  --ranker_svd_dim 32 --ranker_force_alpha 0.95 --ranker_no_pos",
    )
    add_callout(doc, "代码版本说明", "文件名仍是 v62，但当前工作区中的该文件已经包含 v65/v67 特征配置和 180 维原始特征。因此复现 V69 应使用当前这个文件，而不是旧的 Kaggle v41 副本。", kind="caution")

    doc.add_heading("7.3 输出校验", level=2)
    for text in [
        "ZIP 根目录只有 dataset1.csv 与 dataset2.csv，不要嵌套文件夹。",
        "dataset1.csv 恰好 61,051 行；dataset2.csv 恰好 153,420 行。",
        "每行恰好 100 个可解析浮点数，全部有限，无 NaN/Inf。",
        "每行最小值约为 0、最大值约为 1；这验证了行内 min-max 归一化已执行。",
        "固定随机种子后，XGBoost/GPU 版本差异仍可能造成极小浮动；应比较名次与 MRR，而非逐位字节完全相同。",
    ]:
        add_bullet(doc, text, bullet_id)

    doc.add_heading("8. 融合与门控：如何复现后处理", level=1)
    doc.add_heading("8.1 秩融合", level=2)
    add_body(doc, "不同模型的原始分数尺度不可比，因此融合脚本优先把每行分数变成名次。最简单的 reciprocal-rank 融合为：")
    add_equation(doc, "F(d) = w/rank_ours(d) + (1−w)/rank_other(d)")
    add_body(doc, "RRF 使用 1/(k+rank)，Borda 使用线性名次分，rank-product 对名次取加权对数，norm 则直接融合行内归一化分数。dataset1 的实验权重通常偏向自有 V69（0.80～0.90），dataset2 则偏向外部方案（自有权重 0.15～0.25）。")

    doc.add_heading("8.2 置信度门控", level=2)
    add_body(doc, "门控先用 top1 与 top2 的分差除以行内标准差作为置信度，再做稳健中心化和 sigmoid。只有当两模型 top1 不一致时，才按 conf_other−conf_ours 选出最值得切换的 3%、5% 或 8% 行；未选中行给自有模型 0.92 权重，选中行给自有模型 0.06 权重。")
    add_callout(doc, "非自包含依赖", "V71–V94 的“other/open”原始提交在当前目录中没有单独保留；现有混合 ZIP 只保存了结果。若要从头重做这些融合，必须重新取得外部提交。完全可复现的算法主体是 V69 独立方案。", kind="risk")

    doc.add_heading("9. 工程优化与资源估算", level=1)
    add_table(
        doc,
        ["优化点", "实现", "收益"],
        [
            ("一次聚合", "groupby 只在 build_history 执行", "避免每个候选重复扫描历史边"),
            ("压缩键查询", "key=u×base+v，排序数组+searchsorted", "批量 O(log U) 查询，缓存友好"),
            ("节点稠密数组", "dst_count[id]、src_last[id] 等", "节点统计近似 O(1) 索引"),
            ("float32", "计数特征、嵌入、特征矩阵均优先 float32", "内存减半、GPU 训练更快"),
            ("分块构造/预测", "ranker_chunk_rows=1500；输出流式写盘", "控制临时数组峰值，避免一次性预测爆内存"),
            ("稳定排序", "mergesort 处理 time 与 key", "相同时间/分数时结果更可复现"),
            ("GPU→CPU 回退", "捕获训练异常后切换 CPU hist", "牺牲速度但不中断流程"),
            ("特征档位", "v42/v51/v58/v59/v61/v65/v67", "按效果和资源选择，而非始终使用 180 维"),
        ],
        [1900, 3850, 3610],
        font_size=8.7,
        first_col_bold=True,
    )
    add_body(doc, "内存估算：130,000 个排序组 × 100 候选 × 105 特征 × 4 字节约 5.1 GiB；v65 的 167 特征约 8.1 GiB，尚未计入标签、验证矩阵、历史索引和 XGBoost 内部副本。dataset1 的实际验证组少于 130k，因此日志中的峰值会低一些，但仍建议至少 16 GiB 内存；v65 更稳妥是 24 GiB 以上。")
    add_body(doc, "时间复杂度主要来自：历史分组与排序、每个 batch 的 10 个最近位置×正反 transition/skip 查询、SVD 点积 O(B×100×32)，以及树模型训练。feature_batch 已向量化，但 180 维全开仍会造成显著内存压力。")

    doc.add_heading("10. 复现风险、验证边界与改进建议", level=1)
    add_table(
        doc,
        ["问题", "影响", "处理建议"],
        [
            ("验证协议变化", "不同版本 MRR 不完全可比", "只在相同行数、候选模式、树数下做严格消融"),
            ("跨 test 聚合", "可能触及规则边界", "确认规则；必要时关闭 test_* 特征"),
            ("only_dataset 打包", "缺另一个 CSV 时最后报 WinError 2", "同一 output_dir 先生成两文件，再覆盖单数据集"),
            ("外部 open 缺失", "V71–V94 无法从零重建", "以 V69 为自包含基线；单独归档外部提交及哈希"),
            ("模型未保存", "每次预测必须重训", "增加 model.json、feature_profile、参数和依赖版本归档"),
            ("无早停", "固定 850/1200 树可能过拟合或浪费", "保留验证集并启用 early_stopping_rounds"),
            ("时间泄漏检查", "聚合窗口若包含未来边会高估", "所有验证历史必须严格早于验证边；保留 split/时间审计"),
            ("分数非概率", "不能按概率解释阈值", "只用于排序；若需概率需另做校准"),
        ],
        [2000, 3300, 4060],
        font_size=8.6,
        first_col_bold=True,
    )

    doc.add_heading("建议的下一轮优化", level=2)
    add_bullet(doc, "保存训练好的 Booster 与特征清单，拆分 train.py / predict.py，避免每次提交全量重训。", bullet_id)
    add_bullet(doc, "用严格时间切分做多折 backtesting，并让每折都复刻 row 候选分布；用均值和方差而非单次 MRR 选特征。", bullet_id)
    add_bullet(doc, "把图 SVD 替换或补充为时间衰减的 LightGCN/Node2Vec/矩阵分解，但仍以候选行的 ranking loss 训练融合层。", bullet_id)
    add_bullet(doc, "对 dataset1/v65 做 SHAP 或 permutation 消融，减少 167 维中低收益转移特征，降低 1200 树的计算成本。", bullet_id)
    add_bullet(doc, "建立实验注册表：代码哈希、数据哈希、依赖版本、随机种子、验证协议、线上分数与输出 ZIP 哈希同时保存。", bullet_id)

    doc.add_heading("11. 最小可复现伪代码", level=1)
    add_code(
        doc,
        "for dataset in [dataset1, dataset2]:\n"
        "    train = sort_by_time(read_train(dataset))\n"
        "    test  = read_test(dataset)\n"
        "\n"
        "    hist_train, val, fake_test, true_pos = make_validation(\n"
        "        train, test, candidate_mode='row')\n"
        "    history = build_history(hist_train, fake_test, svd_dim=32)\n"
        "\n"
        "    X = feature_batch(history, fake_test)       # group × 100 × F\n"
        "    y = one_positive_per_group(true_pos)\n"
        "    model = XGBRanker(objective='rank:ndcg', eval_metric='ndcg@10')\n"
        "    model.fit(X_train, y_train, group=[100, ...])\n"
        "\n"
        "    alpha = choose_by_validation_mrr(model_score, v18_score)\n"
        "    full_history = build_history(train, test, svd_dim=32)\n"
        "    for test_chunk in chunks(test, 1500):\n"
        "        features = feature_batch(full_history, test_chunk)\n"
        "        score = alpha * row_norm(model.predict(features))\n"
        "              + (1-alpha) * row_norm(v18_score(test_chunk))\n"
        "        write(row_norm(score))\n"
        "\n"
        "zip(dataset1.csv, dataset2.csv)",
    )

    doc.add_heading("12. 源码定位（便于逐段核对）", level=1)
    add_table(
        doc,
        ["主题", "文件与当前行号"],
        [
            ("GCN 归一化与模型", "gcn.py:90–124"),
            ("早停、重训、集成", "gcn.py:152–202, 333–377"),
            ("类别配额分配", "gcn.py:205–273"),
            ("图历史与 SVD", "formal_track1/track1_solution_v62_dataset1_ranker.py:118–415"),
            ("V18 启发式分数", "同上:449–511, 1342–1406"),
            ("特征清单与 profile", "同上:518–929"),
            ("完整 feature_batch", "同上:946–1247"),
            ("伪验证候选", "同上:1264–1339"),
            ("XGBRanker 与 α 融合", "同上:1642–1865"),
            ("全量预测与打包", "同上:1868–2078"),
            ("融合实验", "formal_track1/fuse_submissions.py；fusion_experiments_v71.py"),
            ("门控与按数据集组装", "formal_track1/gated_fusions_v88.py；hybrid_copy_fusions_v83.py"),
        ],
        [2850, 6510],
        font_size=8.8,
        first_col_bold=True,
    )
    add_callout(doc, "最终核对", "本文中的算法、参数和本地分数均来自当前工作区文件与日志；没有把外部网页描述当作实现依据。若源码继续修改，行号和 feature_profile 的实际内容应重新审计。", kind="info")

    set_running_furniture(doc)
    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    build_document()
